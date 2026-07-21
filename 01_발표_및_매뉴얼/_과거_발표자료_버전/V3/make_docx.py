import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_script_docx():
    doc = Document()
    
    # Page setup - Horizontal/Landscape orientation (A4 Size: 11.69 x 8.27 Inches)
    sections = doc.sections
    for section in sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Malgun Gothic'
    style_normal.font.size = Pt(14)  # Increased font size for landscape readability
    style_normal.font.color.rgb = RGBColor(30, 30, 30)
    style_normal.paragraph_format.line_spacing = 1.7
    style_normal.paragraph_format.space_after = Pt(14)
    
    # Custom Heading Style
    style_heading = doc.styles.add_style('SlideHeading', 1) # Paragraph style
    style_heading.font.name = 'Malgun Gothic'
    style_heading.font.size = Pt(20)  # Increased heading size
    style_heading.font.bold = True
    style_heading.font.color.rgb = RGBColor(0, 144, 69) # Green color (#009045)
    style_heading.paragraph_format.space_before = Pt(0)
    style_heading.paragraph_format.space_after = Pt(20)
    
    # Content Data
    slides_data = [
        {
            "num": "Slide 01",
            "title": "타이틀: 바이브코딩으로 실현한 정산 자동화",
            "content": [
                "안녕하십니까, 업무 자동화 사례 발표를 맡은 피플팀 정영욱입니다.",
                "오늘은 개발자가 아닌 실무자가, '바이브코딩'만으로 수일이 걸리던 작업을 단 몇 분으로 줄였는지에 대해 공유해 드리려 합니다."
            ]
        },
        {
            "num": "Slide 02",
            "title": "BEFORE: 수작업의 고통",
            "content": [
                "매월 말, 전도금 정산업무로 인해 심천지사와 유관부서가 겪어야 했던 상황입니다.",
                "[SHIFT 클릭] 지금 보이시는 화면처럼, 100~300장에 달하는 영수증과 증빙을 한 장 한 장 눈으로 읽고, 엑셀에 금액과 날짜, 계정과목을 직접 타이핑해야 했습니다.",
                "하다보면 휴먼에러로 인해 한 글자만 틀려도 오류가 생길 수 있고, 심천지사 뿐 아니라 SCM담당과 재무관리팀 또한 수기 검증해야 합니다. 아래 타임라인과 같이 최소 2~3일은 소요되는 업무였습니다."
            ]
        },
        {
            "num": "Slide 03",
            "title": "문제 정의 & 해결책",
            "content": [
                "좌측의 빨간색 항목이 기존 업무방식의 문제점입니다. 앞서 말씀드린 대로 월말의 업무 병목, 과도한 리드타임, 수동 타이핑, 그리고 반드시 발생할 수밖에 없는 휴먼 에러와 이를 잡아내기 위한 여러 부서의 교차 검증까지 진행됩니다.",
                "저는 해결책이 'AI 바이브코딩'이라고 생각했습니다. 이후 실무자와 충분히 인터뷰를 하고, 직접 AI에게 업무의 규칙을 설명해 나갔습니다. 말하는 대로 설계되고, 실시간으로 작동 방식을 눈으로 확인하며, 업무에 맞춘 시스템을 빠르게 개발할 수 있었습니다."
            ]
        },
        {
            "num": "Slide 04",
            "title": "AI와 대화 시작 (터미널 애니메이션 1)",
            "content": [
                "실제로 제가 AI와 어떻게 대화했는지 보시겠습니다.",
                "[클릭] '심천지사 전도금 정산 업무를 자동화할 건데, 영수증을 읽어서 대계정/소계정을 분류하고 엑셀로 출력해줘. 확인용 웹앱도 만들어줘.'",
                "[클릭] 놀랍게도 단 이 한 단락의 요청만으로 AI는 스스로 5단계 시스템 개발계획을 알려주고 설계하고 코딩을 시작했습니다. 저는 코드 한 줄 직접 짜지 않았습니다."
            ]
        },
        {
            "num": "Slide 05",
            "title": "제작 과정 · 고도화 (터미널 애니메이션 2)",
            "content": [
                "당연히 처음부터 완벽하게 개발될 수는 없었고, 개발 중 발견된 문제들은 바로바로 AI에게 피드백을 주었습니다.",
                "'날짜 형식을 엑셀에 맞게 무조건 통일해 줘', '중복 제출된 영수증은 경고를 띄워 줘'라고 요청하면, AI는 바로 시스템에 반영하여 고도화해 나갔습니다."
            ]
        },
        {
            "num": "Slide 06",
            "title": "📹 대화로 만든 산출물을 보여드립니다",
            "content": [
                "AI와 대화로 만든 결과물을 직접 보여드리겠습니다.",
                "[① 영수증 투입]\n\"영수증 파일들을 폴더에 넣습니다. 기존에는 하나하나 다 출력 후, 수기로 체크해가며 다시 스캔을 했었습니다.\"",
                "[② AI OCR 판독]\n\"정산실행 배치 프로그램을 더블클릭해서, OCR 기능을 통해 자동으로 영수증의 글씨를 판독하고 계정을 분류합니다. 완료된 영수증 파일명은 순번과 계정에 맞추어 순서대로 규칙화 됩니다.\"",
                "[③ 웹앱 검토]\n\"분석이 완료된 후, 웹앱에 접속합니다. 영수증마다 날짜, 계정, 담당자, 금액이 다 반영되었고 재검토가 필요한 항목은 비고란에 표시되어 검토가 빨라집니다.\n기존에는 영수증과 엑셀파일을 열어서 각각 비교하는 번거로움이 있었는데, 한 화면에서 영수증과 결과를 대조해볼 수 있는 편리한 기능까지 추가했습니다. 웹앱에서 수정을 완료하고, 엑셀 최종 저장버튼을 누릅니다.\"",
                "[④ 엑셀 완성]\n\"이제 기존에 쓰던 양식 그대로 데이터가 반영된 정산파일이 생성됩니다.\"\n아직 테스트 단계라 사용하면서 개선할 부분이 조금씩 있으나, 정산업무에 효율성이 기대된다는 실무자의 좋은 피드백이 있었습니다."
            ]
        },
        {
            "num": "Slide 07",
            "title": "시스템 구조 흐름",
            "content": [
                "개발 단계를 정리해보면 크게 6단계입니다.",
                "영수증 투입 ➔ 원클릭 실행 ➔ OCR 판독 ➔ 자동 계정 분류 ➔ 웹앱을 통한 직관적 검토 ➔ 마지막으로 엑셀 자동 완성입니다.",
                "과거에는 최소 2~3일이 걸리던 했던 일이, 이제는 클릭 한 번에 끝나는 구조로 개선이 가능해졌습니다."
            ]
        },
        {
            "num": "Slide 08",
            "title": "결론 및 제언",
            "content": [
                "이번 개발을 통해 해당 업무를 가장 잘 아는 '본인'이 바로 가장 훌륭한 개발자가 된다는 것을 느꼈습니다.",
                "복잡한 코딩을 배우지 않아도 원하는 자동화를 구상할 수 있고, 내가 하는 일의 문제점과 해결 방식만 명확히 알려줄 수 있다면 AI와 함께 구현할 수 있습니다.",
                "오늘 소개해드린 사례가 구성원 분들에게 조금이나마 도움이 되셨으면 좋겠습니다. 우리 구성원 분들께서도 주변의 업무나 일상에서부터 가볍게 시도해보셨으면 좋겠습니다."
            ]
        },
        {
            "num": "Slide 09",
            "title": "마무리",
            "content": [
                "이상으로 발표 마치겠습니다. 감사합니다. 질문이 있으시면 편하게 말씀해주시기 바랍니다."
            ]
        }
    ]

    # Process slides
    for idx, slide in enumerate(slides_data):
        # Header
        h = doc.add_paragraph(style='SlideHeading')
        run_num = h.add_run(f"[{slide['num']}] ")
        run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(0, 144, 69)
        
        run_title = h.add_run(slide['title'])
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(40, 40, 40)
        
        # Table to wrap script for visual structure
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(10.0) # Matches landscape width (11.69 - 1.6 margins)
        
        # Set border / shading for layout look
        cell = table.cell(0,0)
        set_cell_background(cell, "F4F9F5") # Light greenish background for visibility
        cell.width = Inches(10.0)
        
        # Add left border to the cell to look like a callout block
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="009045"/>\n'
            f'  <w:top w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'  <w:bottom w:val="none"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)
        
        # Content inside cell
        cell_p = cell.paragraphs[0]
        cell_p.style = 'Normal'
        cell_p.paragraph_format.space_before = Pt(10)
        cell_p.paragraph_format.space_after = Pt(14)
        cell_p.paragraph_format.line_spacing = 1.7
        
        for p_idx, text in enumerate(slide['content']):
            if p_idx > 0:
                cell_p = cell.add_paragraph()
                cell_p.style = 'Normal'
                cell_p.paragraph_format.space_before = Pt(10)
                cell_p.paragraph_format.space_after = Pt(14)
                cell_p.paragraph_format.line_spacing = 1.7
            
            # Highlight cues like [SHIFT 클릭], [클릭], [① 영수증 투입]
            parts = re.split(r'(\[[^\]\n]+\])', text)
            for part in parts:
                if part.startswith('[') and part.endswith(']'):
                    run = cell_p.add_run(part)
                    run.font.bold = True
                    if '클릭' in part:
                        run.font.color.rgb = RGBColor(220, 80, 80) # Reddish for action cues
                    else:
                        run.font.color.rgb = RGBColor(243, 107, 33) # Orange for steps
                else:
                    run = cell_p.add_run(part)
                    run.font.color.rgb = RGBColor(30, 30, 30)
                    
        # Slide Break except last
        if idx < len(slides_data) - 1:
            doc.add_page_break()
            
    # Save document
    output_path = "c:/Users/20000177/Desktop/Wooktigravity/260630_shenzhen_restart_step_1/_발표자료/V3/발표대본_정영욱.docx"
    doc.save(output_path)
    print(f"Document successfully created at {output_path}")

if __name__ == "__main__":
    create_script_docx()
